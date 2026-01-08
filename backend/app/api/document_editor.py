"""
Document Editor API - Handles document content extraction and conversion
"""
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
import io
import logging
import httpx

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/editor", tags=["Document Editor"])


class DocumentContentRequest(BaseModel):
    storage_url: str
    file_type: str
    file_name: Optional[str] = None  # Optional: original filename for extension-based detection
    document_id: Optional[str] = None  # Optional: if provided, extracted text will be saved
    version_id: Optional[str] = None  # Optional: if provided, saves to document_versions table


class SaveDocumentRequest(BaseModel):
    html_content: str
    file_name: str
    format: str = "docx"  # docx, pdf, txt, html


class DocumentContentResponse(BaseModel):
    content: str
    content_type: str  # html, text
    success: bool
    error: Optional[str] = None
    extracted_text: Optional[str] = None  # Plain text for saving to DB


@router.post("/extract-content", response_model=DocumentContentResponse)
async def extract_document_content(request: DocumentContentRequest):
    """
    Extract content from a document URL and return as HTML/text for editing.
    Supports: PDF, DOCX, DOC, TXT, RTF
    If document_id provided, extracts and saves plain text to database.
    """
    try:
        # Download the file from storage URL
        async with httpx.AsyncClient() as client:
            response = await client.get(request.storage_url, timeout=30.0)
            if response.status_code != 200:
                return DocumentContentResponse(
                    content="",
                    content_type="text",
                    success=False,
                    error=f"Failed to download file: {response.status_code}"
                )
            file_bytes = response.content

        file_type = request.file_type.lower()
        file_name = request.file_name or ""
        content = ""
        content_type = "html"
        extracted_text = ""

        # Extract content based on file type
        if "pdf" in file_type:
            content = await extract_pdf_content(file_bytes)
            extracted_text = content  # For PDF, we already get text
        elif "doc" in file_type or "word" in file_type:
            content = await extract_word_content(file_bytes)
            extracted_text = content
        elif "presentation" in file_type or "powerpoint" in file_type or file_name.lower().endswith(('.pptx', '.ppt')):
            # PowerPoint files
            from app.services.office_processor import office_processor
            extracted_text = office_processor.extract_from_powerpoint(file_bytes)
            content = f"<p>{extracted_text.replace(chr(10)+chr(10), '</p><p>').replace(chr(10), '<br/>')}</p>"
        elif "spreadsheet" in file_type or "excel" in file_type or file_name.lower().endswith(('.xlsx', '.xls', '.csv')):
            # Excel/CSV files
            from app.services.office_processor import office_processor
            extracted_text = office_processor.extract_from_excel(file_bytes, file_name)
            content = f"<pre>{extracted_text}</pre>"
        elif "txt" in file_type or "text" in file_type:
            extracted_text = file_bytes.decode('utf-8', errors='ignore')
            content = f"<p>{extracted_text.replace(chr(10)+chr(10), '</p><p>').replace(chr(10), '<br/>')}</p>"
        elif "rtf" in file_type:
            content = await extract_rtf_content(file_bytes)
            extracted_text = content
        elif "html" in file_type:
            content = file_bytes.decode('utf-8', errors='ignore')
            extracted_text = content
            content_type = "html"
        else:
            # Try to decode as text
            try:
                extracted_text = file_bytes.decode('utf-8', errors='ignore')
                content = f"<p>{extracted_text.replace(chr(10)+chr(10), '</p><p>').replace(chr(10), '<br/>')}</p>"
            except:
                return DocumentContentResponse(
                    content="",
                    content_type="text",
                    success=False,
                    error=f"Unsupported file type: {file_type}"
                )

        # Strip HTML tags from extracted text for database storage
        import re
        plain_text = re.sub('<[^<]+?>', '', extracted_text).strip()

        # Save extracted text to database if document_id provided
        if request.document_id and plain_text:
            try:
                from app.core.supabase_client import supabase
                
                # Update documents table with extracted_text
                doc_update = supabase.table('documents').update({
                    'extracted_text': plain_text
                }).eq('id', request.document_id).execute()
                
                # If version_id provided, also update document_versions
                if request.version_id:
                    version_update = supabase.table('document_versions').update({
                        'content': plain_text
                    }).eq('id', request.version_id).execute()
                    logger.info(f"Saved extracted text to document_versions: {request.version_id}")
                
                logger.info(f"Saved {len(plain_text)} chars of extracted text to document: {request.document_id}")
            except Exception as save_error:
                logger.warning(f"Could not save extracted text to database: {save_error}")

        return DocumentContentResponse(
            content=content,
            content_type=content_type,
            success=True,
            extracted_text=plain_text
        )

    except Exception as e:
        logger.error(f"Error extracting document content: {e}")
        return DocumentContentResponse(
            content="",
            content_type="text",
            success=False,
            error=str(e)
        )


async def extract_pdf_content(file_bytes: bytes) -> str:
    """Extract text from PDF and convert to HTML paragraphs"""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        html_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Use "text" extraction with sort=True for consistent reading order
            # This ensures the same text is extracted regardless of PDF internal structure
            text = page.get_text("text", sort=True)
            
            if text.strip():
                # Convert text to HTML paragraphs
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        # Clean and format
                        clean_para = para.replace('\n', ' ').strip()
                        html_parts.append(f"<p>{clean_para}</p>")
            
            # Add page break marker between pages
            if page_num < len(doc) - 1:
                html_parts.append('<hr/>')
        
        doc.close()
        return ''.join(html_parts) if html_parts else "<p>No text content found in PDF.</p>"
        
    except ImportError:
        logger.warning("PyMuPDF not installed, trying PyPDF2")
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            html_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    paragraphs = text.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            clean_para = para.replace('\n', ' ').strip()
                            html_parts.append(f"<p>{clean_para}</p>")
            
            return ''.join(html_parts) if html_parts else "<p>No text content found in PDF.</p>"
            
        except ImportError:
            return "<p>PDF extraction libraries not available. Please install PyMuPDF or PyPDF2.</p>"
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"<p>Error extracting PDF content: {str(e)}</p>"


def extract_pdf_plain_text(file_bytes: bytes) -> str:
    """Extract plain text from PDF for version comparison (consistent extraction)"""
    try:
        import fitz
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            # Use sort=True for consistent reading order across extractions
            text = page.get_text("text", sort=True)
            if text.strip():
                text_parts.append(text.strip())
        
        doc.close()
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Plain text extraction error: {e}")
        return ""


async def extract_word_content(file_bytes: bytes) -> str:
    """Extract content from Word document and convert to HTML"""
    try:
        import mammoth
        
        result = mammoth.convert_to_html(io.BytesIO(file_bytes))
        html = result.value
        
        # Log any conversion messages
        if result.messages:
            for msg in result.messages:
                logger.debug(f"Mammoth message: {msg}")
        
        return html if html else "<p>No content found in document.</p>"
        
    except ImportError:
        logger.warning("Mammoth not installed, trying python-docx")
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            html_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check for heading styles
                    if para.style.name.startswith('Heading'):
                        level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                        html_parts.append(f"<h{level}>{para.text}</h{level}>")
                    else:
                        html_parts.append(f"<p>{para.text}</p>")
            
            # Extract tables
            for table in doc.tables:
                html_parts.append("<table>")
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        html_parts.append(f"<td>{cell.text}</td>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")
            
            return ''.join(html_parts) if html_parts else "<p>No content found in document.</p>"
            
        except ImportError:
            return "<p>Word document extraction libraries not available. Please install mammoth or python-docx.</p>"
    except Exception as e:
        logger.error(f"Word extraction error: {e}")
        return f"<p>Error extracting Word document: {str(e)}</p>"


async def extract_rtf_content(file_bytes: bytes) -> str:
    """Extract content from RTF document"""
    try:
        from striprtf.striprtf import rtf_to_text
        
        rtf_text = file_bytes.decode('utf-8', errors='ignore')
        plain_text = rtf_to_text(rtf_text)
        
        # Convert to HTML
        paragraphs = plain_text.split('\n\n')
        html_parts = [f"<p>{para.replace(chr(10), '<br/>')}</p>" for para in paragraphs if para.strip()]
        
        return ''.join(html_parts) if html_parts else "<p>No content found in RTF document.</p>"
        
    except ImportError:
        return "<p>RTF extraction library not available. Please install striprtf.</p>"
    except Exception as e:
        logger.error(f"RTF extraction error: {e}")
        return f"<p>Error extracting RTF content: {str(e)}</p>"


@router.post("/convert-to-docx")
async def convert_to_docx(request: SaveDocumentRequest):
    """Convert HTML content to DOCX format for download"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from bs4 import BeautifulSoup
        
        # Parse HTML
        soup = BeautifulSoup(request.html_content, 'html.parser')
        
        # Create Word document
        doc = Document()
        
        # Process HTML elements
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'table']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                para = doc.add_paragraph()
                # Handle formatted text
                for child in element.children:
                    if hasattr(child, 'name'):
                        run = para.add_run(child.get_text())
                        if child.name == 'strong' or child.name == 'b':
                            run.bold = True
                        if child.name == 'em' or child.name == 'i':
                            run.italic = True
                        if child.name == 'u':
                            run.underline = True
                    else:
                        para.add_run(str(child))
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        filename = request.file_name.replace('.docx', '').replace('.doc', '') + '_edited.docx'
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx library not installed")
    except Exception as e:
        logger.error(f"DOCX conversion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/convert-to-pdf")
async def convert_to_pdf(request: SaveDocumentRequest):
    """Convert HTML content to PDF format for download"""
    try:
        from weasyprint import HTML
        
        # Wrap content in proper HTML structure
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; font-size: 12pt; line-height: 1.6; margin: 2cm; }}
                h1 {{ font-size: 24pt; margin-bottom: 12pt; }}
                h2 {{ font-size: 18pt; margin-bottom: 10pt; }}
                h3 {{ font-size: 14pt; margin-bottom: 8pt; }}
                p {{ margin-bottom: 10pt; }}
                table {{ border-collapse: collapse; width: 100%; margin: 10pt 0; }}
                th, td {{ border: 1px solid #333; padding: 8pt; text-align: left; }}
            </style>
        </head>
        <body>
            {request.html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        pdf_bytes = HTML(string=full_html).write_pdf()
        
        filename = request.file_name.replace('.pdf', '').replace('.docx', '').replace('.doc', '') + '_edited.pdf'
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except ImportError:
        raise HTTPException(status_code=500, detail="weasyprint library not installed. Use HTML export instead.")
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/health")
async def editor_health():
    """Check available document processing capabilities"""
    capabilities = {
        "pdf_extract": False,
        "word_extract": False,
        "word_create": False,
        "pdf_create": False,
        "rtf_extract": False
    }
    
    try:
        import fitz
        capabilities["pdf_extract"] = True
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            capabilities["pdf_extract"] = True
        except ImportError:
            pass
    
    try:
        import mammoth
        capabilities["word_extract"] = True
    except ImportError:
        try:
            from docx import Document
            capabilities["word_extract"] = True
        except ImportError:
            pass
    
    try:
        from docx import Document
        capabilities["word_create"] = True
    except ImportError:
        pass
    
    try:
        from weasyprint import HTML
        capabilities["pdf_create"] = True
    except ImportError:
        pass
    
    try:
        from striprtf.striprtf import rtf_to_text
        capabilities["rtf_extract"] = True
    except ImportError:
        pass
    
    try:
        from pdf2docx import Converter
        capabilities["pdf_to_docx"] = True
    except ImportError:
        pass
    
    try:
        from docx2pdf import convert
        capabilities["docx_to_pdf"] = True
    except ImportError:
        pass
    
    return {"status": "ok", "capabilities": capabilities}


# OnlyOffice callback for saving documents
class OnlyOfficeCallback(BaseModel):
    key: str
    status: int
    url: Optional[str] = None
    changesurl: Optional[str] = None
    history: Optional[dict] = None
    users: Optional[list] = None
    actions: Optional[list] = None
    lastsave: Optional[str] = None
    notmodified: Optional[bool] = None
    forcesavetype: Optional[int] = None


async def save_document_to_supabase(document_id: str, file_bytes: bytes, file_name: str):
    """Save the edited document to Supabase as a new version"""
    import os
    from datetime import datetime
    import uuid
    
    # Use the singleton Supabase client for connection pooling
    from app.core.supabase_client import get_supabase_client
    
    supabase = get_supabase_client()
    if not supabase:
        logger.error("Supabase client not available")
        return False
    
    logger.info("Using shared Supabase client (connection pooling)")
    
    try:
        # Get current document info
        doc_result = supabase.table('documents').select('*').eq('id', document_id).single().execute()
        if not doc_result.data:
            logger.error(f"Document not found: {document_id}")
            return False
        
        doc = doc_result.data
        user_id = doc.get('user_id') or doc.get('uploaded_by')
        
        # Determine file extension
        ext = file_name.split('.')[-1].lower() if '.' in file_name else 'docx'
        
        # Generate new storage path for the version
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        new_filename = f"{timestamp}_edited.{ext}"
        storage_path = f"{user_id}/{new_filename}"
        
        # Upload to Supabase Storage
        upload_result = supabase.storage.from_('documents').upload(
            storage_path,
            file_bytes,
            file_options={"content-type": f"application/{'pdf' if ext == 'pdf' else 'vnd.openxmlformats-officedocument.wordprocessingml.document'}"}
        )
        
        # Get signed URL
        url_result = supabase.storage.from_('documents').create_signed_url(storage_path, 3600 * 24 * 365)
        storage_url = url_result.get('signedUrl', '')
        
        # Get current version number
        current_version = doc.get('metadata', {}).get('current_version', 1) if doc.get('metadata') else 1
        new_version = current_version + 1
        
        # Extract text from the new file for documents.extracted_text (used for chatbot embeddings)
        extracted_text = None
        plain_text_for_version = None  # Separate plain text for version comparison
        try:
            file_type = doc.get('file_type', '').lower()
            if 'pdf' in file_type:
                extracted_text = await extract_pdf_content(file_bytes)
                plain_text_for_version = extract_pdf_plain_text(file_bytes)  # Consistent extraction for versions
            elif 'doc' in file_type or 'word' in file_type:
                extracted_text = await extract_word_content(file_bytes)
            elif 'rtf' in file_type:
                extracted_text = await extract_rtf_content(file_bytes)
            else:
                try:
                    extracted_text = file_bytes.decode('utf-8', errors='ignore')
                except:
                    pass
            
            # Strip HTML tags for plain text storage in documents table
            if extracted_text:
                import re
                extracted_text = re.sub('<[^<]+?>', '', extracted_text).strip()
                logger.info(f"Extracted {len(extracted_text)} chars from edited document")
            
            # Use plain_text_for_version if available, otherwise use extracted_text
            if not plain_text_for_version:
                plain_text_for_version = extracted_text
                
        except Exception as extract_error:
            logger.warning(f"Could not extract text from edited document: {extract_error}")
        
        # Update document metadata with new version
        updated_metadata = doc.get('metadata') or {}
        updated_metadata['current_version'] = new_version
        updated_metadata['last_edited_at'] = datetime.now().isoformat()
        updated_metadata['versions'] = updated_metadata.get('versions', [])
        updated_metadata['versions'].append({
            'version': new_version,
            'storage_path': storage_path,
            'storage_url': storage_url,
            'created_at': datetime.now().isoformat(),
            'file_size': len(file_bytes),
            'source': 'onlyoffice_edit'
        })
        
        # Build update data - always update extracted_text with latest version for chatbot
        update_data = {
            'storage_path': storage_path,
            'file_size': len(file_bytes),
            'metadata': updated_metadata,
            'updated_at': datetime.now().isoformat()
        }
        if extracted_text:
            update_data['extracted_text'] = extracted_text
        
        # Update the document record
        update_result = supabase.table('documents').update(update_data).eq('id', document_id).execute()
        
        # Also insert into document_versions table for version comparison feature
        try:
            version_record = {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'version_number': new_version,
                'content': plain_text_for_version or extracted_text or storage_path,  # Use plain text for consistent comparison
                'change_summary': f'Edited via OnlyOffice (v{new_version})',
                'created_by': user_id,
                'major_version': new_version,
                'minor_version': 0,
            }
            
            version_result = supabase.table('document_versions').insert(version_record).execute()
            
            if version_result.data:
                logger.info(f"Created document_versions record for v{new_version} with {'plain text' if plain_text_for_version else 'extracted text' if extracted_text else 'storage path'}")
            else:
                logger.warning(f"Failed to create document_versions record: {version_result}")
        except Exception as version_error:
            # Log but don't fail - the main document save was successful
            logger.warning(f"Could not create document_versions record: {version_error}")
        
        if update_result.data:
            logger.info(f"Document {document_id} saved as v{new_version}")
            return True
        else:
            logger.error(f"Failed to update document {document_id}: {update_result}")
            return False
        
    except Exception as e:
        logger.error(f"Error saving to Supabase: {e}")
        return False



@router.post("/onlyoffice-callback")
async def onlyoffice_callback(callback: OnlyOfficeCallback):
    """
    Handle OnlyOffice Document Server callbacks.
    Status codes:
    - 0: No document with this key in the cache
    - 1: Document being edited
    - 2: Document ready for saving
    - 3: Document saving error
    - 4: Document closed with no changes
    - 6: Document being edited, but the document version was changed
    - 7: Error while force saving the document
    """
    logger.info(f"OnlyOffice callback received: status={callback.status}, key={callback.key}")
    
    # Extract document ID from key (format: documentId_timestamp)
    document_id = callback.key.split('_')[0] if '_' in callback.key else callback.key
    
    if callback.status == 2:  # Document ready for saving
        if callback.url:
            try:
                # Download the saved document from OnlyOffice
                async with httpx.AsyncClient() as client:
                    response = await client.get(callback.url, timeout=60.0)
                    if response.status_code == 200:
                        file_bytes = response.content
                        logger.info(f"Downloaded edited document: {len(file_bytes)} bytes")
                        
                        # Determine file type from URL or default to docx
                        file_ext = 'docx'
                        if callback.url.endswith('.pdf'):
                            file_ext = 'pdf'
                        
                        # Save to Supabase as new version
                        success = await save_document_to_supabase(
                            document_id,
                            file_bytes,
                            f"edited.{file_ext}"
                        )
                        
                        if success:
                            logger.info(f"Document {document_id} saved successfully as new version")
                        else:
                            logger.error(f"Failed to save document {document_id} to Supabase")
                    else:
                        logger.error(f"Failed to download saved document: {response.status_code}")
            except Exception as e:
                logger.error(f"Error saving document: {e}")
        
        return {"error": 0}  # Tell OnlyOffice we handled it
    
    elif callback.status == 4:  # Closed with no changes
        logger.info(f"Document closed with no changes: {callback.key}")
        return {"error": 0}
    
    elif callback.status == 6:  # Force save
        if callback.url:
            logger.info(f"Force save triggered for: {callback.key}")
            try:
                async with httpx.AsyncClient() as client:
                    response = await client.get(callback.url, timeout=60.0)
                    if response.status_code == 200:
                        file_bytes = response.content
                        file_ext = 'docx'
                        if callback.url.endswith('.pdf'):
                            file_ext = 'pdf'
                        
                        success = await save_document_to_supabase(
                            document_id,
                            file_bytes,
                            f"edited.{file_ext}"
                        )
                        logger.info(f"Force save {'successful' if success else 'failed'} for {document_id}")
            except Exception as e:
                logger.error(f"Error during force save: {e}")
        return {"error": 0}
    
    # For other statuses, just acknowledge
    return {"error": 0}


class PdfToDocxRequest(BaseModel):
    storage_url: str
    filename: Optional[str] = None


@router.post("/pdf-to-docx")
async def convert_pdf_to_docx(request: PdfToDocxRequest):
    """
    Convert PDF to DOCX format for editing.
    Uses multiple methods for best quality:
    1. LibreOffice (best quality, preserves images and layout)
    2. PyMuPDF + python-docx (preserves images)
    3. pdf2docx (fallback)
    Returns the DOCX file as a download.
    """
    import tempfile
    import os
    import subprocess
    
    # Download the PDF file
    async with httpx.AsyncClient() as client:
        response = await client.get(request.storage_url, timeout=120.0)
        if response.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Failed to download PDF: {response.status_code}")
        pdf_bytes = response.content
    
    # Create temp files for conversion
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
        pdf_temp.write(pdf_bytes)
        pdf_path = pdf_temp.name
    
    docx_path = pdf_path.replace('.pdf', '.docx')
    conversion_success = False
    conversion_method = "none"
    
    try:
        # Method 1: Try LibreOffice first (best quality, preserves images and layout)
        logger.info(f"Attempting LibreOffice PDF to DOCX conversion: {pdf_path}")
        try:
            # Find LibreOffice executable on Windows
            soffice_paths = [
                'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
                'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
                'soffice',  # Linux/Mac PATH
                '/usr/bin/soffice',
                '/usr/bin/libreoffice',
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            ]
            
            soffice_cmd = None
            for path in soffice_paths:
                if os.path.exists(path):
                    soffice_cmd = path
                    logger.info(f"Found LibreOffice at: {path}")
                    break
            
            if soffice_cmd:
                out_dir = os.path.dirname(pdf_path)
                # Use writer_pdf_import filter for better PDF import quality
                result = subprocess.run([
                    soffice_cmd, '--headless', '--infilter=writer_pdf_import',
                    '--convert-to', 'docx:MS Word 2007 XML',
                    '--outdir', out_dir,
                    pdf_path
                ], capture_output=True, timeout=120)
                
                if result.returncode == 0 and os.path.exists(docx_path):
                    logger.info("LibreOffice conversion successful")
                    conversion_success = True
                    conversion_method = "LibreOffice"
                else:
                    logger.warning(f"LibreOffice conversion failed: {result.stderr.decode() if result.stderr else 'No error message'}")
            else:
                logger.warning("LibreOffice not found on system")
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            logger.warning(f"LibreOffice conversion error: {e}")
        
        # Method 2: Try PyMuPDF + python-docx (preserves images)
        if not conversion_success:
            logger.info("Trying PyMuPDF + python-docx conversion")
            try:
                import fitz  # PyMuPDF
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Open PDF with PyMuPDF
                pdf_doc = fitz.open(pdf_path)
                doc = Document()
                
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    
                    # Extract images from the page
                    image_list = page.get_images()
                    images_added = set()
                    
                    for img_index, img_info in enumerate(image_list):
                        try:
                            xref = img_info[0]
                            if xref in images_added:
                                continue
                            images_added.add(xref)
                            
                            # Extract image
                            pix = fitz.Pixmap(pdf_doc, xref)
                            
                            # Convert to RGB if necessary
                            if pix.n - pix.alpha > 3:
                                pix = fitz.Pixmap(fitz.csRGB, pix)
                            
                            # Save to bytes
                            img_bytes = pix.tobytes("png")
                            
                            # Add to document
                            img_stream = io.BytesIO(img_bytes)
                            try:
                                # Add image with max width of 6 inches
                                paragraph = doc.add_paragraph()
                                run = paragraph.add_run()
                                run.add_picture(img_stream, width=Inches(6))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            except Exception as img_add_err:
                                logger.warning(f"Could not add image: {img_add_err}")
                            
                        except Exception as img_err:
                            logger.warning(f"Error extracting image {img_index}: {img_err}")
                    
                    # Extract text blocks with their positions
                    blocks = page.get_text("dict")["blocks"]
                    
                    for block in blocks:
                        if block["type"] == 0:  # Text block
                            for line in block.get("lines", []):
                                line_text = ""
                                for span in line.get("spans", []):
                                    line_text += span.get("text", "")
                                
                                if line_text.strip():
                                    para = doc.add_paragraph()
                                    run = para.add_run(line_text)
                                    
                                    # Try to preserve font size from first span
                                    if line.get("spans"):
                                        font_size = line["spans"][0].get("size", 11)
                                        run.font.size = Pt(font_size)
                    
                    # Add page break between pages (except last)
                    if page_num < len(pdf_doc) - 1:
                        doc.add_page_break()
                
                pdf_doc.close()
                doc.save(docx_path)
                
                if os.path.exists(docx_path):
                    conversion_success = True
                    conversion_method = "PyMuPDF"
                    logger.info("PyMuPDF + python-docx conversion successful")
                    
            except ImportError as ie:
                logger.warning(f"PyMuPDF/python-docx not available: {ie}")
            except Exception as e:
                logger.warning(f"PyMuPDF conversion failed: {e}")
        
        # Method 3: Fallback to pdf2docx
        if not conversion_success:
            logger.info("Falling back to pdf2docx for conversion")
            try:
                from pdf2docx import Converter
                cv = Converter(pdf_path)
                cv.convert(docx_path)
                cv.close()
                conversion_success = True
                conversion_method = "pdf2docx"
                logger.info("pdf2docx conversion successful")
            except ImportError:
                logger.error("pdf2docx not installed")
            except Exception as pdf2docx_err:
                logger.error(f"pdf2docx conversion failed: {pdf2docx_err}")
        
        if not conversion_success or not os.path.exists(docx_path):
            raise HTTPException(
                status_code=500, 
                detail="PDF to DOCX conversion failed. Please install LibreOffice for best results. Download from: https://www.libreoffice.org/"
            )
        
        # Read the DOCX file
        with open(docx_path, 'rb') as f:
            docx_bytes = f.read()
        
        # Determine filename
        filename = request.filename or "converted.docx"
        if not filename.endswith('.docx'):
            filename = filename.rsplit('.', 1)[0] + '.docx'
        
        logger.info(f"PDF to DOCX conversion successful using {conversion_method}: {len(docx_bytes)} bytes")
        
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-Conversion-Method": conversion_method
            }
        )
        
    finally:
        # Cleanup temp files
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)


class DocxToPdfRequest(BaseModel):
    filename: Optional[str] = None


@router.post("/docx-to-pdf")
async def convert_docx_to_pdf(file: UploadFile = File(...)):
    """
    Convert DOCX to PDF format.
    Accepts a DOCX file upload and returns a PDF.
    """
    try:
        import tempfile
        import os
        import subprocess
        
        # Read the uploaded DOCX file
        docx_bytes = await file.read()
        
        # Create temp files for conversion
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            docx_temp.write(docx_bytes)
            docx_path = docx_temp.name
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Try using LibreOffice for conversion (works on Linux/Windows/Mac)
            # This is more reliable than docx2pdf which requires MS Word
            try:
                # Try LibreOffice first
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(docx_path),
                    docx_path
                ], capture_output=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception("LibreOffice conversion failed")
                    
            except (FileNotFoundError, Exception) as e:
                logger.warning(f"LibreOffice not available: {e}, trying docx2pdf")
                # Fallback to docx2pdf (requires MS Word on Windows)
                try:
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                except ImportError:
                    # Use python-docx and reportlab as last resort
                    from docx import Document
                    from reportlab.lib.pagesizes import letter
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.units import inch
                    
                    doc = Document(docx_path)
                    c = canvas.Canvas(pdf_path, pagesize=letter)
                    width, height = letter
                    
                    y = height - inch
                    for para in doc.paragraphs:
                        if y < inch:
                            c.showPage()
                            y = height - inch
                        
                        # Simple text wrapping
                        text = para.text
                        if text:
                            # Set font size based on style
                            font_size = 12
                            if para.style.name.startswith('Heading'):
                                font_size = 16 if '1' in para.style.name else 14
                            
                            c.setFont("Helvetica", font_size)
                            
                            # Wrap text to fit page width
                            max_width = width - 2 * inch
                            words = text.split()
                            line = ""
                            for word in words:
                                test_line = f"{line} {word}".strip()
                                if c.stringWidth(test_line, "Helvetica", font_size) < max_width:
                                    line = test_line
                                else:
                                    if line:
                                        c.drawString(inch, y, line)
                                        y -= font_size + 4
                                        if y < inch:
                                            c.showPage()
                                            y = height - inch
                                    line = word
                            
                            if line:
                                c.drawString(inch, y, line)
                                y -= font_size + 8
                    
                    c.save()
            
            # Read the PDF file
            if not os.path.exists(pdf_path):
                raise HTTPException(status_code=500, detail="PDF conversion failed - output file not created")
                
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
            
            # Determine filename
            filename = file.filename or "document.pdf"
            if not filename.endswith('.pdf'):
                filename = filename.rsplit('.', 1)[0] + '.pdf'
            
            logger.info(f"DOCX to PDF conversion successful: {len(pdf_bytes)} bytes")
            
            return StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'}
            )
            
        finally:
            # Cleanup temp files
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
                
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX to PDF conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
